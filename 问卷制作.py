import pandas as pd
import shutil
from pathlib import Path
from tqdm import tqdm
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 需要安装 python-docx 包
# pip install python-docx

def modify_docx(file_path, member_name_list):
    doc = Document(file_path)
    # 获取表格
    table = doc.tables[-1]  # 假设最后一页只有一个表格

    # 填写表格
    row = table.rows[0]  # 表格的第一行
    row.cells[1].text = member_name_list[0]  # 填写第二列
    row.cells[2].text = member_name_list[1]  # 填写第三列
    row.cells[3].text = member_name_list[2]  # 填写第四列
    # 修改字体格式和段落格式
    for cell in row.cells:
        # 修改字体大小和名称
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)  # 字体大小
                run.font.name = '楷体'  # 字体名称
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')#强制替换中文字体
            # 修改段落对齐方式
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平居中对齐
                            
    # 保存修改后的文档
    doc.save(file_path)
    
# -- config ---
# Path是一个包不用管
root_path = Path("T3团队问卷")
member_doc_path = Path("data\External Leaders-T3合并.docx")
leader_doc_path = Path("data\External Team Members-T3合并.docx")

if __name__ == "__main__":
    # 删除第一行, 第二行作为列名
    data = pd.read_excel(Path("data\人工智能团队行为项目-0427.xlsx"))
    data.columns = data.iloc[0].values
    data = data.drop(0)

    # 清空文件夹
    shutil.rmtree(root_path, ignore_errors=True)
    leader_num = data["团队主管"].nunique()
    for i in tqdm(range(1, leader_num+1)):
        # step 1: 新建文件夹
        root_path.joinpath(f"团队{i}").mkdir(parents=True, exist_ok=True)
        
        # step 2: 成员文件copy
        member_name_list = data.iloc[i-1].values[-3:]
        for member_name in member_name_list:
            shutil.copy(leader_doc_path, root_path.joinpath(f"团队{i}").joinpath(f"团队{i}成员-{member_name}.docx"))
        
        # step 3: 主管文件copy
        leader_name = data.iloc[i-1]["团队主管"]
        shutil.copy(member_doc_path, root_path.joinpath(f"团队{i}").joinpath(f"团队{i}主管-{leader_name}.docx"))
        
        # 如果主管的问卷中不需要加入成员信息，可以注释掉下面这行
        modify_docx(root_path.joinpath(f"团队{i}").joinpath(f"团队{i}主管-{leader_name}.docx"), member_name_list)
